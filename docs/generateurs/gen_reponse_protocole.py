# -*- coding: utf-8 -*-
"""Réponses au protocole de tests v6.0 -> Word (.docx) soigné + export PDF."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r'c:\Users\Jorda\familystore-pos'
LOGO = os.path.join(ROOT, 'frontend', 'src', 'assets', 'logo-fs.jpg')
OUT_DOCX = os.path.join(ROOT, 'REPONSE_PROTOCOLE_TESTS_V6.docx')

WINE      = RGBColor(0x7A, 0x1D, 0x2E)
WINE_DARK = RGBColor(0x4A, 0x0E, 0x1C)
GOLD      = RGBColor(0xB8, 0x86, 0x0B)
GREY      = RGBColor(0x66, 0x66, 0x66)
INK       = RGBColor(0x2B, 0x2B, 0x2B)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREEN     = RGBColor(0x1D, 0x7A, 0x4E)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'; normal.font.size = Pt(10.5); normal.font.color.rgb = INK


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hex_color)
    tcPr.append(sh)


def no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'), 'nil'); borders.append(e)
    tblPr.append(borders)


# ── Bannière ──────────────────────────────────────────────────────────────────
banner = doc.add_table(rows=1, cols=2)
banner.alignment = WD_TABLE_ALIGNMENT.CENTER
no_borders(banner)
banner.columns[0].width = Inches(1.6)
banner.columns[1].width = Inches(5.4)
c_logo, c_txt = banner.rows[0].cells
shade(c_logo, 'FDF9F0'); shade(c_txt, '7A1D2E')

lp = c_logo.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(LOGO):
    try:
        lp.add_run().add_picture(LOGO, width=Inches(1.45))
    except Exception:
        r = lp.add_run('FS'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = GOLD

tp = c_txt.paragraphs[0]
tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = tp.add_run('  RÉPONSES AU PROTOCOLE DE TESTS v6.0')
r1.bold = True; r1.font.size = Pt(17); r1.font.color.rgb = WHITE
tp2 = c_txt.add_paragraph()
r2 = tp2.add_run('  Family Store POS — 12 juillet 2026')
r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0xE8, 0xC4, 0xA0)

doc.add_paragraph()
intro = doc.add_paragraph(
    "Suite à votre protocole de tests du 08.07.2026, l'ensemble des recommandations a été traité. "
    "Les modifications sont déjà déployées : fermez puis rouvrez l'application (ou Ctrl+F5) pour en profiter."
)
intro.runs[0].italic = True; intro.runs[0].font.color.rgb = GREY


def section(titre, statut='CORRIGÉ'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    run = p.add_run('  ' + titre + '   ')
    run.bold = True; run.font.size = Pt(12.5); run.font.color.rgb = WINE
    if statut:
        rs = p.add_run(f'[{statut}]')
        rs.bold = True; rs.font.size = Pt(9.5)
        rs.font.color.rgb = GREEN if statut in ('CORRIGÉ', 'FAIT', 'AMÉLIORÉ') else GOLD
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2'); bottom.set(qn('w:color'), 'D9CEB9')
    pbdr.append(bottom)
    pPr.append(pbdr)


def para(texte, gras=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if gras:
        rg = p.add_run(gras); rg.bold = True; rg.font.color.rgb = WINE_DARK
    p.add_run(texte)


def puce(gras, texte=''):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if gras:
        rg = p.add_run(gras); rg.bold = True; rg.font.color.rgb = WINE_DARK
    if texte:
        p.add_run(texte)


# ── Contenu ──────────────────────────────────────────────────────────────────
section('REC#1 — Curseur de la souris sur les éléments non cliquables')
para("Dans les rapports, le curseur reste désormais une flèche standard sur tout ce qui n'est pas cliquable. "
     "Seuls les vrais boutons (onglets, exports, colonnes triables) affichent la main.")

section('REC#2 — Nomenclature des noms de produits')
para("Les noms de produits s'affichent maintenant selon la nomenclature dans tous les rapports "
     "(« vanilla » devient « Vanilla », « balea serum… » devient « Balea Serum… »), y compris dans l'export PDF. "
     "Les unités (ml, g…) sont préservées telles quelles.")

section('REC#3 — Référence du pourcentage calculé', 'RÉPONSE + AMÉLIORÉ')
para("Chaque pourcentage compare la période en cours à la ", gras='Réponse : ')
p = doc.paragraphs[-1]; p.add_run('période précédente équivalente').bold = True
p.add_run(" : semaine en cours ↔ semaine dernière, mois en cours ↔ mois dernier, année en cours ↔ année dernière. "
          "Formule : (valeur actuelle − valeur précédente) ÷ valeur précédente × 100.")
para("la valeur de référence est désormais affichée sous chaque pourcentage "
     "(ex. « ↑ 14,6 % — vs 127 137 XAF (période préc.) »). Quand la période précédente est vide, "
     "la mention « Pas de référence » l'indique clairement.", gras='Représentation améliorée : ')

section('REC#4 — Différence entre les deux graphiques CA', 'RÉPONSE + CORRIGÉ')
para("dans la version testée, les deux graphiques montraient effectivement la même donnée sous deux formes. "
     "Ils ont été différenciés pour apporter deux lectures :", gras='Réponse : ')
puce('CA par période (barres) : ', 'combien chaque jour / période a rapporté.')
puce('CA cumulé (courbe) : ', "la progression du chiffre d'affaires additionné au fil de la période, jusqu'au total. "
     "On y lit le rythme des ventes et si le mois avance bien.")

section('REC#5 — Filtres sur les colonnes du « Journal des ventes par produit »', 'FAIT')
para("Chaque colonne se trie d'un clic sur son en-tête (▲ croissant / ▼ décroissant) : Produit, Qté vendue, "
     "CA généré, Nb transactions, Prix moyen. Un champ de recherche permet aussi de filtrer par nom de produit.")

section('REC#6 — Offre marketing modifiable sur la facture (import/export Excel)', 'FAIT')
para("Nouvelle rubrique « Paramètres magasin → Offre marketing (facture) » avec les 5 champs demandés : "
     "TITRE_OFFRE, MESSAGE_OFFRE, VALIDITE_OFFRE, CALL_TO_ACTION, SALUTATION_FIN.")
puce('Deux façons de modifier : ', "directement à l'écran, ou Exporter (CSV) → édition dans Excel → Importer (CSV).")
puce('Votre fichier OFFRE_MARKETING.csv ', "s'importe tel quel (format « CLE: texte »).")
puce('Mise en gras : ', "les mots entre astérisques *comme ceci* sont bien imprimés en gras sur le ticket.")
puce('Champ vide ', "= ligne non imprimée sur le ticket.")
puce('Accents : ', "l'export produit un fichier qui s'ouvre sans caractères bizarres dans Excel.")

section("REC#7 — Données du diagramme d'affluence horaire", 'RÉPONSE')
para("ce sont les ventes réelles de la période affichée (le mois sélectionné). Chaque vente est classée "
     "dans sa case jour de la semaine × heure de la journée d'après son horodatage de caisse. L'intensité de la "
     "couleur représente le montant vendu (CA) dans ce créneau, en proportion du créneau le plus fort. "
     "Le diagramme montre donc quand le magasin encaisse le plus.", gras='Réponse : ')

section("REC#8 — Filtres sur l'Historique des factures", 'FAIT')
para("La page dispose maintenant des mêmes filtres rapides que le Journal des ventes : Aujourd'hui · Cette semaine · "
     "Ce mois · Plage de dates · Tout, plus un sélecteur par caissière. Le filtrage s'applique instantanément et "
     "reste exact même sur plusieurs pages de factures.")

section('REC#8 (bis) — Logo Family Store', 'FAIT')
para("Le logo officiel Family Store (couronne et ornements dorés) est intégré en tête du menu d'administration.")

doc.add_paragraph()
note = doc.add_paragraph()
rn = note.add_run("Merci pour la qualité de ce protocole de tests — n'hésitez pas à nous transmettre la prochaine série de retours.")
rn.italic = True; rn.font.color.rgb = GREY

# ── En-tête & pied de page ───────────────────────────────────────────────────
hdr = sec.header.paragraphs[0]
hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rh = hdr.add_run('Family Store · Réponses protocole v6.0')
rh.font.size = Pt(8); rh.font.color.rgb = GREY

ftr = sec.footer.paragraphs[0]
ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = ftr.add_run('Family Store POS · Bonamoussadi, Douala     —     Page ')
rf.font.size = Pt(8); rf.font.color.rgb = GREY
fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE')
ftr._p.append(fld1)

doc.save(OUT_DOCX)
print('DOCX OK ->', OUT_DOCX)

# Export PDF via Word en forçant une imprimante aux métriques correctes.
# (L'imprimante par défaut « Generic / Text Only » fausse les largeurs de
#  caractères → espaces parasites dans le PDF.)
try:
    import win32com.client
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        word.ActivePrinter = 'Microsoft Print to PDF'
        d = word.Documents.Open(OUT_DOCX, ReadOnly=True)
        d.ExportAsFixedFormat(OutputFileName=OUT_DOCX.replace('.docx', '.pdf'), ExportFormat=17)
        d.Close(False)
    finally:
        word.Quit()
    print('PDF OK ->', OUT_DOCX.replace('.docx', '.pdf'))
except Exception as e:
    print('PDF non genere:', e)
