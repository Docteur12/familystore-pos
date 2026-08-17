# -*- coding: utf-8 -*-
"""Recap client « Grosse mise a jour » -> Word (.docx) soigne + export PDF."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r'c:\Users\Jorda\familystore-pos'
LOGO = os.path.join(ROOT, 'frontend', 'public', 'apple-touch-icon.png')
OUT_DOCX = os.path.join(ROOT, 'RECAP_GROSSE_MISE_A_JOUR.docx')

WINE      = RGBColor(0x7A, 0x1D, 0x2E)
WINE_DARK = RGBColor(0x4A, 0x0E, 0x1C)
GOLD      = RGBColor(0xB8, 0x86, 0x0B)
GREY      = RGBColor(0x66, 0x66, 0x66)
INK       = RGBColor(0x2B, 0x2B, 0x2B)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Marges
sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hex_color)
    tcPr.append(sh)


def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'nil')
        borders.append(e)
    tblPr.append(borders)


# ── Bannière titre (tableau 1 cellule, fond bordeaux) ────────────────────────
banner = doc.add_table(rows=1, cols=2)
banner.alignment = WD_TABLE_ALIGNMENT.CENTER
no_borders(banner)
banner.columns[0].width = Inches(1.0)
banner.columns[1].width = Inches(6.0)
c_logo, c_txt = banner.rows[0].cells
shade(c_logo, '7A1D2E')
shade(c_txt, '7A1D2E')

# Logo
lp = c_logo.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(LOGO):
    try:
        lp.add_run().add_picture(LOGO, width=Inches(0.7))
    except Exception:
        r = lp.add_run('FS'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = GOLD

# Titre dans la bannière
tp = c_txt.paragraphs[0]
tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = tp.add_run('GROSSE MISE A JOUR')
r1.bold = True; r1.font.size = Pt(22); r1.font.color.rgb = WHITE
tp2 = c_txt.add_paragraph()
r2 = tp2.add_run('Family Store — 2 juin 2026')
r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(0xE8, 0xC4, 0xA0)

doc.add_paragraph()
intro = doc.add_paragraph(
    "Une mise a jour importante a ete deployee. Voici, en langage simple, "
    "tout ce qui change et s'ameliore pour vous et vos equipes."
)
intro.runs[0].italic = True
intro.runs[0].font.color.rgb = GREY


def section(titre):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('  ' + titre)
    run.bold = True
    run.font.size = Pt(13.5)
    run.font.color.rgb = WINE
    # liseré bas
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2'); bottom.set(qn('w:color'), 'D9CEB9')
    pbdr.append(bottom)
    pPr.append(pbdr)


def puce(gras, texte=''):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if gras:
        rg = p.add_run(gras); rg.bold = True; rg.font.color.rgb = WINE_DARK
    if texte:
        p.add_run(texte)


# ── Contenu ──────────────────────────────────────────────────────────────────
section('Caisse (caissiers)')
puce('Article divers : ', "vendre un produit pas encore enregistre — il suffit de saisir son prix, la vente n'est plus bloquee. Le gestionnaire le regularise ensuite.")
puce('Scan plus fiable : ', "fini les « produit introuvable » aleatoires ; le code est lu en entier, et le champ se vide apres une erreur.")
puce('Clavier numerique ', "sur les champs de scan (tablette) pour eviter les caracteres parasites.")
puce('Audit admin : ', "voir les suppressions de vente faites par l'administrateur (qui / quand), propres a votre caisse.")

section('Ticket de caisse')
puce('Plus de TVA ', "affichee sur le ticket.")
puce('Noms de produits longs ', "raccourcis proprement (…), et presentation plus aeree.")
puce('Adresse simplifiee : ', "« Bonamoussadi · Douala ».")
puce('Nouvelle offre fidelite : ', "si le client n'a eu aucune reduction, le ticket affiche « Merci pour votre achat — Family Store vous offre 5 % de reduction sur votre prochain achat. Presentez cette facture a la caisse. » Lorsqu'une reduction est appliquee, le message de remerciement habituel revient.")

section('Magazinier')
puce('Creation produit : ', "ajoute automatiquement a la reception (avec sa quantite), formulaire qui se vide pour enchainer — plus besoin de re-scanner.")
puce('Code-barres ', "qui se reinitialise apres validation d'une reception.")
puce("Messages d'erreur clairs ", "(ex. « code-barres deja utilise »).")
puce('Recherche ', "d'un produit dans le tableau et l'historique ; heure affichee en plus de la date.")
puce("Prix d'achat et de vente : ", "le magazinier peut les renseigner (creation, ligne de reception, tableau). Une fois fixes, ils sont verrouilles — le gestionnaire ne peut plus les modifier.")

section('Gestionnaire de stock')
puce('Nouvelle page « Articles divers » : ', "liste des ventes d'articles non referencees a regulariser, avec creation du produit pre-remplie (nom + prix).")
puce('Prix verrouilles : ', "quand le magazinier a fixe un prix, le gestionnaire le voit pre-rempli et non modifiable, avec la mention « Prix ajoute/modifie par … le … ».")

section('Administrateur')
puce('Supprimer une vente ', "(Journal des ventes) : le stock est restaure automatiquement et l'action est tracee.")
puce('Supprimer une facture ', "+ fin des doublons de factures a l'impression.")
puce("Valeur de l'entrepot : ", "un total + une valeur par produit (quantite entrepot × prix d'achat).")
puce('TVA retiree ', "partout (comptabilite, exports, parametres, factures).")
puce("Navigation entre espaces sans se reconnecter : ", "raccourci « Changer d'espace » (Caisse / Gestion de stock / Magazinier) et bouton « Retour admin ». Les passages de l'administrateur sont traces dans l'audit.")
puce('Depannage en caisse : ', "l'administrateur peut encaisser ; la vente est clairement marquee « Depannage ».")

section('Connexion & securite')
puce('Sessions plus longues (30 jours) : ', "moins de reconnexions sur les appareils toujours allumes.")
puce('Reconnexion automatique : ', "en cas de session expiree, retour a la page de connexion avec le message « Session expiree ».")

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(14)
rn = note.add_run(
    "Toutes les actions sensibles (suppressions, modifications de prix, acces aux espaces) "
    "sont enregistrees dans le journal d'audit pour une tracabilite complete."
)
rn.italic = True; rn.font.color.rgb = GREY


# ── En-tete & pied de page ───────────────────────────────────────────────────
hdr = sec.header.paragraphs[0]
hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rh = hdr.add_run('Family Store · Mise a jour')
rh.font.size = Pt(8); rh.font.color.rgb = GREY

ftr = sec.footer.paragraphs[0]
ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = ftr.add_run('Family Store POS · Bonamoussadi, Douala · Tel : 682 263 435     —     Page ')
rf.font.size = Pt(8); rf.font.color.rgb = GREY
# champ numero de page
fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE')
ftr._p.append(fld1)

doc.save(OUT_DOCX)
print('DOCX OK ->', OUT_DOCX)

# ── Export PDF (via Word) ────────────────────────────────────────────────────
try:
    from docx2pdf import convert
    convert(OUT_DOCX, OUT_DOCX.replace('.docx', '.pdf'))
    print('PDF OK ->', OUT_DOCX.replace('.docx', '.pdf'))
except Exception as e:
    print('PDF non genere:', e)
